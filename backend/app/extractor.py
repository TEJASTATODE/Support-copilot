"""
Document text extraction — PDF and DOCX.

Why these two formats:
- PDF: most common format for policies, manuals, reports
- DOCX: most common format for internal docs, SOPs, guides
- Both are extractable without OCR (digital text, not scanned images)

Pipeline: file bytes → extract text → existing chunking → embedding → pgvector
The extractor is just the first step — everything after is unchanged.
"""
import io
from pathlib import Path


def extract_pdf(file_bytes: bytes) -> str:
    """
    Extract text from a PDF using PyMuPDF (fitz).
    Handles multi-page PDFs, preserves paragraph structure.
    """
    import fitz  # PyMuPDF

    text_parts = []
    with fitz.open(stream=file_bytes, filetype="pdf") as doc:
        for page_num, page in enumerate(doc):
            text = page.get_text("text")
            if text.strip():
                text_parts.append(f"--- Page {page_num + 1} ---\n{text.strip()}")

    full_text = "\n\n".join(text_parts)

    if not full_text.strip():
        raise ValueError(
            "No extractable text found in PDF. "
            "This may be a scanned document — please copy-paste the text manually."
        )

    return full_text


def extract_docx(file_bytes: bytes) -> str:
    """
    Extract text from a DOCX file using python-docx.
    Preserves paragraph structure, extracts table text too.
    """
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                parts.append(row_text)

    full_text = "\n\n".join(parts)

    if not full_text.strip():
        raise ValueError("No text found in DOCX file.")

    return full_text


def extract_text(filename: str, file_bytes: bytes) -> str:
    """
    Route to the right extractor based on file extension.
    Returns extracted plain text ready for chunking.
    """
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        return extract_pdf(file_bytes)
    elif ext in (".docx", ".doc"):
        return extract_docx(file_bytes)
    elif ext == ".txt":
        return file_bytes.decode("utf-8", errors="ignore")
    else:
        raise ValueError(
            f"Unsupported file type '{ext}'. "
            f"Supported: .pdf, .docx, .txt"
        )